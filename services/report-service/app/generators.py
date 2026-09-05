import io
from datetime import datetime, timezone, timedelta
_LOCAL_TZ = timezone(timedelta(hours=2))


# ── PDF (reportlab) ───────────────────────────────────────────────────────

def generate_pdf(data: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    font_dir = os.path.join(os.path.dirname(__import__("reportlab").__file__), "fonts")
    pdfmetrics.registerFont(TTFont("Vera", os.path.join(font_dir, "Vera.ttf")))
    pdfmetrics.registerFont(TTFont("Vera-Bold", os.path.join(font_dir, "VeraBd.ttf")))

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    for s in styles.byName.values():
        s.fontName = "Vera"
    title_style = ParagraphStyle("Title", parent=styles["Heading1"], fontSize=16, fontName="Vera-Bold")
    story = []

    m = data["meeting"]

    # zaglavlje: rukovodilac
    story.append(Paragraph(f"<b>Rukovodilac sastanka:</b> {data['organizer_name']}", styles["Normal"]))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph("ZAPISNIK SA SASTANKA", title_style))
    story.append(Spacer(1, 0.5*cm))

    # osnovni podaci
    info = [
        ["Tema:", m.topic],
        ["Kategorija:", data["category_name"]],
        ["Organizaciona celina:", data["org_unit_name"]],
        ["Tip:", m.meeting_type],
        ["Zakazano:", m.scheduled_at.strftime("%d.%m.%Y %H:%M")],
        ["Mesto:", f"{m.location}, {m.room}"],
        ["Status:", m.status],
    ]
    t = Table(info, colWidths=[5*cm, 11*cm])
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Vera-Bold"),
        ("FONTNAME", (1, 0), (1, -1), "Vera"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(t)
    story.append(Spacer(1, 0.5*cm))

    # uvodna reč
    if m.intro:
        story.append(Paragraph("<b>Uvodna reč:</b>", styles["Normal"]))
        story.append(Paragraph(m.intro, styles["Normal"]))
        story.append(Spacer(1, 0.3*cm))

    # dnevni red sa predlozima
    story.append(Paragraph("<b>Dnevni red:</b>", styles["Normal"]))
    for entry in data["agenda"]:
        item = entry["item"]
        story.append(Paragraph(f"{item.order_no}. {item.title}", styles["Normal"]))
        if item.discussion:
            story.append(Paragraph(f"<i>Diskusija:</i> {item.discussion}", styles["Normal"]))
        for pr in entry["proposals"]:
            story.append(Paragraph(
                f"    <b>{pr['ucesnik']}:</b> {pr['sadrzaj']}",
                styles["Normal"]
            ))
    story.append(Spacer(1, 0.3*cm))

    # učesnici
    story.append(Paragraph("<b>Učesnici:</b>", styles["Normal"]))
    part_data = [["Ime", "Uloga", "Organizacija", "Planiran", "Prisustvo"]]
    for p in data["participants"]:
        part_data.append([p["ime"], p["uloga"], p["organizacija"], p["planiran"], p["prisustvo"]])
    pt = Table(part_data, colWidths=[4*cm, 3*cm, 4*cm, 2*cm, 3*cm])
    pt.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Vera-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Vera"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#ecf0f1")]),
    ]))
    story.append(pt)
    story.append(Spacer(1, 0.3*cm))

    # zaključak
    if m.conclusion:
        story.append(Paragraph("<b>Zaključak:</b>", styles["Normal"]))
        story.append(Paragraph(m.conclusion, styles["Normal"]))

    # podnožje: zapisničar
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph(
        f"<b>Zapisničar:</b> {data['recorder_name']}",
        styles["Normal"]
    ))
    story.append(Paragraph(
        f"<i>Generisano: {datetime.now(_LOCAL_TZ).strftime('%d.%m.%Y %H:%M')}</i>",
        styles["Normal"]
    ))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


# ── XLSX (openpyxl) ───────────────────────────────────────────────────────

def generate_xlsx(data: dict) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = Workbook()
    ws = wb.active
    ws.title = "Izvestaj prisustva"

    m = data["meeting"]
    header_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    bold = Font(bold=True)

    # zaglavlje
    ws["A1"] = "IZVEŠTAJ O PRISUSTVU"
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells("A1:E1")

    rows_info = [
        ("Tema:", m.topic),
        ("Zakazano:", m.scheduled_at.strftime("%d.%m.%Y %H:%M")),
        ("Rukovodilac:", data["organizer_name"]),
        ("Zapisničar:", data["recorder_name"]),
        ("Status:", m.status),
    ]
    r = 3
    for label, val in rows_info:
        ws.cell(row=r, column=1, value=label).font = bold
        ws.cell(row=r, column=2, value=val)
        r += 1

    # tabela učesnika
    r += 1
    headers = ["Ime", "Uloga", "Organizacija", "Planiran", "Prisustvo"]
    for c, h in enumerate(headers, start=1):
        cell = ws.cell(row=r, column=c, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    r += 1
    for p in data["participants"]:
        ws.cell(row=r, column=1, value=p["ime"])
        ws.cell(row=r, column=2, value=p["uloga"])
        ws.cell(row=r, column=3, value=p["organizacija"])
        ws.cell(row=r, column=4, value=p["planiran"])
        ws.cell(row=r, column=5, value=p["prisustvo"])
        r += 1

    for col, width in zip("ABCDE", [25, 15, 20, 12, 12]):
        ws.column_dimensions[col].width = width

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ── DOCX (python-docx) ────────────────────────────────────────────────────

def generate_docx(data: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    m = data["meeting"]

    # zaglavlje: rukovodilac
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_head.add_run(f"Rukovodilac sastanka: {data['organizer_name']}")
    run.bold = True

    title = doc.add_heading("ZAPISNIK SA SASTANKA", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # osnovni podaci
    info = [
        ("Tema", m.topic),
        ("Kategorija", data["category_name"]),
        ("Organizaciona celina", data["org_unit_name"]),
        ("Tip", m.meeting_type),
        ("Zakazano", m.scheduled_at.strftime("%d.%m.%Y %H:%M")),
        ("Mesto", f"{m.location}, {m.room}"),
        ("Status", m.status),
    ]
    tbl = doc.add_table(rows=len(info), cols=2)
    tbl.style = "Light Grid Accent 1"
    for i, (label, val) in enumerate(info):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = str(val)

    if m.intro:
        doc.add_heading("Uvodna reč", level=1)
        doc.add_paragraph(m.intro)

    # dnevni red sa predlozima
    doc.add_heading("Dnevni red", level=1)
    for entry in data["agenda"]:
        item = entry["item"]
        doc.add_paragraph(f"{item.order_no}. {item.title}", style="List Number")
        if item.discussion:
            p = doc.add_paragraph()
            run = p.add_run(f"Diskusija: {item.discussion}")
            run.italic = True
        for pr in entry["proposals"]:
            p = doc.add_paragraph()
            run_name = p.add_run(f"{pr['ucesnik']}: ")
            run_name.bold = True
            p.add_run(pr['sadrzaj'])

    # učesnici
    doc.add_heading("Učesnici", level=1)
    pt = doc.add_table(rows=1, cols=5)
    pt.style = "Light Grid Accent 1"
    hdr = pt.rows[0].cells
    for i, h in enumerate(["Ime", "Uloga", "Organizacija", "Planiran", "Prisustvo"]):
        hdr[i].text = h
    for p in data["participants"]:
        row = pt.add_row().cells
        row[0].text = p["ime"]
        row[1].text = p["uloga"]
        row[2].text = p["organizacija"]
        row[3].text = p["planiran"]
        row[4].text = p["prisustvo"]

    if m.conclusion:
        doc.add_heading("Zaključak", level=1)
        doc.add_paragraph(m.conclusion)

    # podnožje: zapisničar
    doc.add_paragraph("")
    p_footer = doc.add_paragraph()
    run = p_footer.add_run(f"Zapisničar: {data['recorder_name']}")
    run.bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()